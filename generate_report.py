
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

def create_duoc_report():
    doc = Document()

    # --- Configuración de Estilos ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    # Márgenes: Superior/Izquierdo 4 cm, Inferior/Derecho 2.5 cm
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(4)
        section.left_margin = Cm(4)
        section.bottom_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # --- PORTADA ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DUOC UC\nESCUELA DE INFORMÁTICA Y TELECOMUNICACIONES\n")
    run.bold = True
    
    doc.add_paragraph("\n" * 5)
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("INFORME TÉCNICO: PLATAFORMA DE GESTIÓN DE INCENDIOS FORESTALES Y URBANOS - MUNICIPALIDAD VALLE DEL SOL")
    run.bold = True
    run.size = Pt(14)

    doc.add_paragraph("\n" * 8)

    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    info.add_run("Integrantes: Ignacio Salazar, Javier Quiroga\n")
    info.add_run("Carrera: Ingeniería en Informática\n")
    info.add_run("Docente: [Nombre del Docente]\n")
    info.add_run("Fecha: 09 de Abril, 2026")

    doc.add_page_break()

    # --- TABLA DE CONTENIDO ---
    doc.add_heading('Tabla de Contenidos', level=1)
    doc.add_paragraph("(Generada automáticamente por Word)")
    doc.add_page_break()

    # --- RESUMEN ---
    doc.add_heading('Resumen', level=1)
    doc.add_paragraph(
        "El presente proyecto detalla la arquitectura y diseño de una plataforma tecnológica para la Municipalidad Valle del Sol, "
        "enfocada en la detección, monitoreo y alerta temprana de incendios. Utilizando un enfoque de microservicios "
        "serverless sobre AWS y Supabase, se busca optimizar la respuesta ante emergencias y mejorar la comunicación con la ciudadanía."
    )
    doc.add_paragraph("Palabras clave: Microservicios, AWS, Supabase, Incendios, Gestión de Emergencias.")
    
    doc.add_heading('Abstract', level=1)
    doc.add_paragraph(
        "This project details the architecture and design of a technological platform for the Valle del Sol Municipality, "
        "focused on the detection, monitoring, and early warning of fires. Using a serverless microservices approach "
        "on AWS and Supabase, the goal is to optimize emergency response and improve communication with the community."
    )
    doc.add_paragraph("Keywords: Microservices, AWS, Supabase, Fires, Emergency Management.")
    doc.add_page_break()

    # --- 1. PLANTEAMIENTO DEL PROBLEMA ---
    doc.add_heading('1. Planteamiento del Problema', level=1)
    doc.add_paragraph(
        "La Municipalidad Valle del Sol enfrenta una amenaza constante de incendios forestales y urbanos. "
        "Actualmente, la gestión de estas emergencias es ineficiente debido a la fragmentación de la información, "
        "la dependencia de procesos manuales y la falta de herramientas de visualización en tiempo real. "
        "Esto resulta en reportes imprecisos, tiempos de respuesta lentos y una comunicación deficiente con los ciudadanos."
    )

    # --- 2. JUSTIFICACIÓN ---
    doc.add_heading('2. Justificación', level=1)
    doc.add_paragraph(
        "La implementación de esta plataforma es vital para reducir el impacto de las catástrofes. "
        "Se justifica mediante el uso de tecnologías de vanguardia que permiten escalabilidad y resiliencia sin altos costos operacionales."
    )
    
    doc.add_heading('Justificación de Tecnologías y Servicios', level=2)
    techs = [
        ("Vercel & Next.js", "Provee un frontend unificado y responsivo con despliegue continuo (CI/CD) eficiente."),
        ("AWS API Gateway", "Punto de entrada único que gestiona la seguridad y el enrutamiento de peticiones."),
        ("AWS Lambda", "Permite ejecutar microservicios de forma independiente y escalable, pagando solo por el uso."),
        ("AWS Cognito", "Gestiona la autenticación y autorización mediante JWT de forma segura y escalable."),
        ("Supabase (PostgreSQL + RLS)", "Base de datos robusta con Seguridad a Nivel de Fila (RLS) y capacidades en tiempo real para el mapa."),
        ("Amazon S3", "Almacenamiento seguro para evidencias multimedia (fotos/videos) reportadas por ciudadanos."),
        ("Amazon SQS/SNS", "Garantizan la entrega de alertas y notificaciones masivas de forma asíncrona."),
        ("AWS X-Ray", "Provee observabilidad total sobre la traza de los microservicios, facilitando la depuración.")
    ]
    for tech, desc in techs:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(f"{tech}: ")
        run.bold = True
        p.add_run(desc)

    # --- 3. ESTADO DEL ARTE ---
    doc.add_heading('3. Estado del Arte / Situación Actual', level=1)
    doc.add_paragraph(
        "Los sistemas actuales son aislados. No existe una integración entre el reporte ciudadano y el despacho de brigadas. "
        "La literatura actual sugiere que las arquitecturas basadas en microservicios y eventos son las más adecuadas para sistemas de misión crítica."
    )

    # --- 5. OBJETIVOS ---
    doc.add_heading('5. Objetivos', level=1)
    doc.add_heading('Objetivo General', level=2)
    doc.add_paragraph("Desarrollar una plataforma integral basada en microservicios para la gestión eficiente de incendios en la Municipalidad Valle del Sol.")
    
    doc.add_heading('Objetivos Específicos', level=2)
    objs = [
        "Implementar un módulo de reporte ciudadano con soporte multimedia y GPS.",
        "Diseñar un dashboard de monitoreo geográfico en tiempo real.",
        "Establecer un sistema de alertas masivas automatizado.",
        "Garantizar la resiliencia del sistema mediante patrones como Circuit Breaker."
    ]
    for obj in objs:
        doc.add_paragraph(obj, style='List Bullet')

    # --- 7. RESULTADOS Y PRODUCTOS ---
    doc.add_heading('7. Resultados y productos esperados', level=1)
    doc.add_paragraph("A continuación se presenta el diagrama de arquitectura propuesto:")
    
    # Insertar Imagen del Diagrama
    img_path = 'Diagrama sin título.drawio.png'
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Cm(15))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("Figura 1: Arquitectura de Microservicios Free Tier - Valle del Sol").alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        doc.add_paragraph("[Imagen del diagrama no encontrada en el directorio]")

    # --- 11. REFERENCIAS ---
    doc.add_heading('11. Referencias bibliográficas', level=1)
    doc.add_paragraph("Duoc UC. (2026). Formato de Documentos Académicos. Biblioteca Duoc UC.")
    doc.add_paragraph("Amazon Web Services. (2026). AWS Lambda Documentation.")
    doc.add_paragraph("Supabase. (2026). Row Level Security Guide.")

    # Guardar documento
    doc_name = "Informe_Final_Valle_Del_Sol.docx"
    doc.save(doc_name)
    print(f"Documento '{doc_name}' creado exitosamente.")

if __name__ == "__main__":
    create_duoc_report()
